import re
import pandas as pd
import glob
import os
import docxpy
from nltk import tokenize
import docx

# creating lists to store data
balance_sheet = []
income_statements = []
notes_numbers = []
final_list_files = []
final_list_sentences_count = []

# intializing input and output folders
path = "D:\\MSIS\\RA\\FS2018docx-tag-20230309T072019Z-001\\Tagged\\"
out_path = "D:\\MSIS\\RA\\FS2018docx-tag-20230309T072019Z-001\\output\\"


# removes dates and values equal to zero contained in the list
def clean_list(l):
    if len(l) != 0:
        res = [i.replace(" ", "") for i in l if
               i not in ['0', '0.0', '0.00', '00.00', '00.0', '000', '00', '$0', '$0.0', '$0.00', '$00.00', '$00.0',
                         '$000', '$00', '', '-', ' ', '2010', '2011', '2012', '2013', '2014', '2015', '2016', '2017',
                         '2018', '2019', '2020', '2021']]
        return res
    else:
        return []


# removes empty values from list
def rmv_bad_elements(l):
    l1 = [i for i in l if i != ""]
    return l1


# removes dates of all formats from list
def rmv_dates(txt):
    regEx1 = '\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|(?:\d{1,2}[-/th|st|nd|rd\s]*)?(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z\s,.]*(?:\d{1,2}[-/th|st|nd|rd)\s,]*)?(?:\d{2,4})|\d{4} to present|prior to \d{4}|in \d{4}|In \d{4}|until \d{4}|for \d{4}|of \d{4}|from the \d{4}|from \d{4}'
    txt = re.sub(regEx1, '', txt)
    return txt


# returns a list of all numbers(only digits)
def find_values(l):
    l1 = []
    for i in l:
        bs = re.findall(r'(?:\d{0,3}\,)*\d{1,3}[\.[0-9]+]?|\$\s*[\d+,\d+]+[\.[0-9]+]?', i)
        clean = clean_list(bs)
        if len(clean) > 0:
            l1.append(clean[0])
    return l1


# returns a list of all numbers(both words and digits)
def find_values_txt(text):
    regex = r'\$\s*[\d+,\d+]+[\.[0-9]+]?|$[0-9]|\b(?:\d+\.\d+|\d+)\b'
    numbers_in_words = r'(?i)\b(twenty|thirty|forty|fifty|sixty|seventy|eighty|ninety|hundred|thousand)\s*(one|two|three|four|five|six|seven|eight|nine)?\b|\b(zero|one|two|three|four|five|six|seven|eight|nine|ten|eleven|twelve|thirteen|fourteen|fifteen|sixteen|seventeen|eighteen|nineteen)\b'
    nlist = re.findall(regex, text)
    wlist = []
    matches = re.findall(numbers_in_words, text)
    for i in matches:
        x = list(i)
        w = " ".join(x).strip()
        wlist.append(w)
    f = nlist + wlist
    return f
    # notes_numbers.append(len(f))


# counts the numbers of sentences containing more than 3 words
def sentence_counter(l):
    for i in l:
        if len(i.split(" ")) < 3:
            l.remove(i)
    # for i in l:
    #     print(i)
    #     print('__________________')
    final_list_sentences_count.append(len(l))


# removes sentences from list if it contains less than 3 words
def list_cleaner(l):
    for i in l:
        if len(i.split(" ")) < 3:
            l.remove(i)


# creates a docx file and writes all sentences in the file
def write_docx(l, filename):
    doc = docx.Document()
    doc.add_heading(filename, 0)
    for i in l:
        i = i.replace('\n', ' ')
        i = i.replace('\t', ' ')
        doc.add_paragraph(i)
    doc.save(out_path + filename + ".docx")


def text_split(text, filename):
    l1 = text.split("(-B-)")  # splits sections based on tags
    l2 = text.split("(-I-)")
    l3 = text.split("(-N-)")

    l = tokenize.sent_tokenize(l3[1])  # creates a list of sentences using tokenize function
    list_cleaner(l)
    sentence_counter(l)
    write_docx(l, filename)
    notes_txt = rmv_dates(l3[1])
    notes_val = find_values_txt(notes_txt)

    bs_txt = rmv_dates(l1[1])
    ins_txt = rmv_dates(l2[1])
    bs_list = bs_txt.split('\n')
    ins_list = ins_txt.split('\n')
    bs_list = rmv_bad_elements(bs_list)  # calls all necessary functions and creates a list values from
    ins_list = rmv_bad_elements(ins_list)  # balance sheet and income statement
    final_bs = find_values(bs_list)
    final_ins = find_values(ins_list)

    balance_sheet.append(len(final_bs))
    income_statements.append(len(final_ins))
    notes_numbers.append(len(notes_val))


# extracts text from docx file
def getText(file):
    text = docxpy.process(file)
    return text


try:
    for i, doc in enumerate(glob.iglob(path + "*.docx")):
        filename = doc.split('\\')[-1][:-5]
        path1 = os.path.abspath(doc)
        txt = getText(path1)
        # print(filename)
        text_split(txt, filename)
        final_list_files.append(filename)
except:
    print('Error ' + filename)

print(len(final_list_files), len(balance_sheet), len(income_statements), len(notes_numbers),
      len(final_list_sentences_count))

# creates a dataframe to store all the lists
df = pd.DataFrame.from_dict(
    {'Filename': final_list_files, 'Balance Sheet count': balance_sheet, 'Income Statement count': income_statements,
     'Notes numbers': notes_numbers, 'Notes sentences': final_list_sentences_count})

# creates an excel file and stores all data in dataframe
df.to_excel("D:\\MSIS\\RA\\FS2018docx-tag-20230309T072019Z-001\\count.xlsx", header=True, index=False)