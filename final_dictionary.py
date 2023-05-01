import re
from nltk import tokenize
import pandas as pd
# nltk.download('punkt')
import glob
import os
import docxpy
import docx

# initialize lists to store data
final_list_files = []
final_list_sentences_count = []
final_list_word_count = []
final_list_number_count = []
total_sentences_count = []
total_word_count = []
total_number_count = []

path = ''                # path to input folder
out_path = ''           # path to output folder
filename = ''

# dictionary containing all the words
d = {"goal": [], "will": [], "expect": [], "we expect": [], "and expect": [], "but expect": [], "do not expect": [],
     "company expects": [], "corporation expects": [], "firm expects": [], "management expects": [], "and expects": [],
     "but expects": [],
     "does not expect": [], "is expected": [], "are expected": [], "not expected": [], "is expecting": [],
     "are expecting": [],
     "not expecting": [], "normally expect": [], "normally expects": [], "currently expect": [],
     "currently expects": [],
     "also expect": [], "and  also expects": [], "we aim": [], "and aim": [], "but aim": [], "do not aim": [],
     "company aims": [],
     "corporation aims": [], "firm aims": [], "management aims": [], "and aims": [], "but aims": [], "does not aim": [],
     "is aimed": [],
     "are aimed": [], "not aimed": [], "is aiming": [], "are aiming": [], "not aiming": [], "normally aim": [],
     "normally aims": [],
     "currently aim": [], "currently aims": [], "also aim": [], "also aims": [], "anticipate": [], "we anticipate": [],
     "and anticipate": [],
     "but aticipate": [], "do not anticipate": [], "company anticipates": [], "corporation anticipates": [],
     "firm anticipates": [],
     "management anticipates": [], "and anticipates": [], "but anticipates": [], "does not anticipate": [],
     "is anticipated": [],
     "are anticipated": [], "not anticipated": [], "is anticipating": [], "are anticipating": [],
     "not anticipating": [],
     "normally anticipate": [], "normally anticipates": [], "currently anticipate": [], "currently anticipates": [],
     "also anticipate": [],
     "also anticipates": [], "assume": [], "we assume": [], "and assume": [], "but assume": [], "do not assume": [],
     "company  assumes": [],
     "corporation  assumes": [], "firm  assumes": [], "management  assumes": [], "and  assumes": [], "but  assumes": [],
     "does not  assume": [],
     "is assumed": [], "are assumed": [], "not assumed": [], "is assuming": [], "are assuming": [], "not assuming": [],
     "normally assume": [], "normally  assumes": [],
     "currently assume": [], "currently  assumes": [], "also assume": [], "also  assumes": [], "commit": [],
     "we commit": [], "and commit": [],
     "but commit": [], "do not commit": [], "company commits": [], "corporation commits": [], "firm commits": [],
     "management commits": [],
     "and commits": [], "but commits": [], "does not commit": [], "is committed": [], "are committed": [],
     "not committed": [], "is committing": [], "are committing": [],
     "not committing": [], "normally commit": [], "normally commits": [], "currently commit": [],
     "currently commits": [], "also commit": [], "also commits": [],
     "we estimate": [], "and estimate": [], "but estimate": [], "do not estimate": [], "company estimates": [],
     "corporation estimates": [],
     "firm estimates": [], "management estimates": [], "and estimates": [], "but estimates": [],
     "does not estimate": [], "is estimated": [],
     "are estimated": [], "not estimated": [], "is estimating": [], "are estimating": [], "not estimating": [],
     "normally estimate": [],
     "normally estimates": [], "currently estimate": [], "currently estimates": [], "also estimate": [],
     "also estimates": [], "we forecast": [],
     "and forecast": [], "but forecast": [], "do not forecast": [], "company forecasts": [],
     "corporation forecasts": [], "firm forecasts": [],
     "management forecasts": [], "and forecasts": [], "but forecasts": [], "does not forecast": [], "is forecasted": [],
     "are forecasted": [],
     "not forecasted": [], "is forecasting": [], "are forecasting": [], "not forecasting": [], "normally forecast": [],
     "normally forecasts": [],
     "currently forecast": [], "currently forecasts": [], "also forecast": [], "also forecasts": [], "foresee": [],
     "we foresee": [], "and foresee": [],
     "but foresee": [], "do not foresee": [], "company foresees": [], "corporation foresees": [], "firm foresees": [],
     "management foresees": [],
     "and foresees": [], "but foresees": [], "does not foresee": [], "is foreseen": [], "are foreseen": [],
     "not foreseen": [], "is foreseeing": [],
     "are foreseeing": [], "not foreseeing": [], "normally foresee": [], "normally foresees": [],
     "currently foresee": [], "currently foresees": [],
     "also foresee": [], "also foresees": [], "we hope": [], "and hope": [], "but hope": [], "do not hope": [],
     "company hopes": [], "corporation hopes": [],
     "firm hopes": [], "management hopes": [], "and hopes": [], "but hopes": [], "does not hope": [], "is hoped": [],
     "are hoped": [], "not hoped": [], "is hoping": [],
     "are hoping": [], "not hoping": [], "normally hope": [], "normally hopes": [], "currently hope": [],
     "currently hopes": [], "also hope": [], "also hopes": [],
     "intend": [], "we intend": [], "and intend": [], "but intend": [], "do not intend": [], "company intends": [],
     "corporation intends": [], "firm intends": [],
     "management intends": [], "and intends": [], "but intends": [], "does not intend": [], "is intended": [],
     "are intended": [], "not intended": [],
     "is intending": [], "are intending": [], "not intending": [], "normally intend": [], "normally intends": [],
     "currently intend": [],
     "currently intends": [], "also intend": [], "also intends": [], "we plan": [], "and plan": [], "but plan": [],
     "do not plan": [],
     "company plans": [], "corporation plans": [], "firm plans": [], "management plans": [], "and plans": [],
     "but plans": [], "does not plan": [],
     "is planed": [], "are planed": [], "not planed": [], "is planning": [], "are planning": [], "not planning": [],
     "normally plan": [], "normally plans": [],
     "currently plan": [], "currently plans": [], "also plan": [], "also plans": [], "we project": [],
     "and project": [], "but project": [], "do not project": [],
     "company projects": [], "corporation projects": [], "firm projects": [], "management projects": [],
     "and projects": [], "but projects": [],
     "does not project": [], "is projected": [], "are projected": [], "not projected": [], "is projecting": [],
     "are projecting": [],
     "not projecting": [], "normally project": [], "normally projects": [], "currently project": [],
     "currently projects": [], "also project": [],
     "also projects": [], "seek": [], "we seek": [], "and seek": [], "but seek": [], "do not seek": [],
     "company seeks": [], "corporation seeks": [],
     "firm seeks": [], "management seeks": [], "and seeks": [], "but seeks": [], "does not seek": [], "is sought": [],
     "are sought": [], "not sought": [],
     "is seeking": [], "are seeking": [], "not seeking": [], "normally seek": [], "normally seeks": [],
     "currently seek": [], "currently seeks": [],
     "also seek": [], "also seeks": [], "we target": [], "and target": [], "but target": [], "do not target": [],
     "company targets": [],
     "corporation targets": [], "firm targets": [], "management targets": [], "and targets": [], "but targets": [],
     "does not target": [],
     "is targeted": [], "are targeted": [], "not targeted": [], "is targeting": [], "are targeting": [],
     "not targeting": [],
     "normally target": [], "normally targets": [], "currently target": [], "currently targets": [], "also target": [],
     "also targets": [],
     "believe": [], "we believe": [], "and believe": [], "but believe": [], "do not believe": [],
     "company believes": [],
     "corporation believes": [], "firm believes": [], "management believes": [], "and believes": [], "but believes": [],
     "does not believe": [], "is believed": [], "are believed": [], "not believed": [], "is believing": [],
     "are believing": [],
     "not believing": [], "normally believe": [], "normally believes": [], "currently believe": [],
     "currently believes": [],
     "also believe": [], "also believes": [], "objective": [], "company objective": [], "corporation objective": [],
     "firm objective": [],
     "management objective": [], "also expects": [], "also will": [], "and will": [], "anticipates": [],
     "are willing": [], "assumes": [],
     "believes": [], "but will": [], "commits": [], "currently aimed": [], "currently aiming": [],
     "currently anticipated": [],
     "currently anticipating": [], "currently assumed": [], "currently assuming": [], "currently believed": [],
     "currently believing": [],
     "currently committed": [], "currently committing": [], "currently estimated": [], "currently estimating": [],
     "currently expected": [],
     "currently expecting": [], "currently forecasted": [], "currently forecasting": [], "currently foreseeing": [],
     "currently foreseen": [], "currently hoped": [], "currently hoping": [], "currently intended": [],
     "currently intending": [],
     "currently planed": [], "currently planning": [], "currently projected": [], "currently projecting": [],
     "currently seeking": [],
     "currently sought": [], "currently targeted": [], "currently targeting": [], "currently will": [],
     "currently willing": [],
     "do not will": [], "does not will": [], "expects": [], "foresees": [], "intends": [], "is willing": [],
     "normally will": [],
     "not willing": [], "now aim": [], "now aimed": [], "now aiming": [], "now aims": [], "now anticipate": [],
     "now anticipated": [],
     "now anticipates": [], "now anticipating": [], "now assume": [], "now assumed": [], "now assumes": [],
     "now assuming": [],
     "now believe": [], "now believed": [], "now believes": [], "now believing": [], "now commit": [],
     "now commits": [],
     "now committed": [], "now committing": [], "now estimate": [], "now estimated": [], "now estimates": [],
     "now estimating": [],
     "now expect": [], "now expected": [], "now expecting": [], "now expects": [], "now forecast": [],
     "now forecasted": [],
     "now forecasting": [], "now forecasts": [], "now foresee": [], "now foreseeing": [], "now foreseen": [],
     "now foresees": [],
     "now hope": [], "now hoped": [], "now hopes": [], "now hoping": [], "now intend": [], "now intended": [],
     "now intending": [],
     "now intends": [], "now plan": [], "now planed": [], "now planning": [], "now plans": [], "now project": [],
     "now projected": [],
     "now projecting": [], "now projects": [], "now seek": [], "now seeking": [], "now seeks": [], "now sought": [],
     "now target": [],
     "now targeted": [], "now targeting": [], "now targets": [], "now will": [], "now willing": [], "seeks": [],
     "still aim": [],
     "still aimed": [], "still aiming": [], "still aims": [], "still anticipate": [], "still anticipated": [],
     "still anticipates": [],
     "still anticipating": [], "still assume": [], "still assumed": [], "still assumes": [], "still assuming": [],
     "still believe": [],
     "still believed": [], "still believes": [], "still believing": [], "still commit": [], "still commits": [],
     "still committed": [],
     "still committing": [], "still estimate": [], "still estimated": [], "still estimates": [], "still estimating": [],
     "still expect": [],
     "still expected": [], "still expecting": [], "still expects": [], "still forecast": [], "still forecasted": [],
     "still forecasting": [],
     "still forecasts": [], "still foresee": [], "still foreseeing": [], "still foreseen": [], "still foresees": [],
     "still hope": [],
     "still hoped": [], "still hopes": [], "still hoping": [], "still intend": [], "still intended": [],
     "still intending": [],
     "still intends": [], "still plan": [], "still planed": [], "still planning": [], "still plans": [],
     "still project": [],
     "still projected": [], "still projecting": [], "still projects": [], "still seek": [], "still seeking": [],
     "still seeks": [],
     "still sought": [], "still target": [], "still targeted": [], "still targeting": [], "still targets": [],
     "still will": [],
     "still willing": [], "will be": [], "we are going": []}

# remove unnecessary period and return clean text
def cleaner(txt):
    txt = txt.replace('etc.', 'etc')
    txt = txt.replace('approx.', 'approx')
    txt = txt.replace('mins.', 'mins')
    return txt

# count numbers(both in digits and in words)
def number_counter(text):
    # regex is regular expression to detect numbers in digits
    regex = r'\$\s*[\d+,\d+]+[\.[0-9]+]?|$[0-9]|\b(?:\d+\.\d+|\d+)\b'
    # numbers_in_words is regular expression to detect number in words
    numbers_in_words = r'(?i)\b(twenty|thirty|forty|fifty|sixty|seventy|eighty|ninety|hundred|thousand)\s*(one|two|three|four|five|six|seven|eight|nine)?\b|\b(zero|one|two|three|four|five|six|seven|eight|nine|ten|eleven|twelve|thirteen|fourteen|fifteen|sixteen|seventeen|eighteen|nineteen)\b'
    nlist = re.findall(regex, text)                        # findall creates a list of all values that matches regex
    wlist = []
    matches = re.findall(numbers_in_words, text)
    for i in matches:
        x = list(i)
        w = " ".join(x).strip()
        wlist.append(w)
    f = nlist + wlist
    total_number_count.append(len(f))
    l = [i for i in f if i in list(d.keys())]
    final_list_number_count.append(len(l))                  # adds total number count to final_list_number_count[]


def sentence_counter(l):
    l1 = []
    for i in l:
        if len(i.split(" ")) < 3:
            l.remove(i)
    total_sentences_count.append(len(l))              # count total number of sentences having 3 or more words
    for i in l:                                       # count sentences containing word from dictionary
        for j in list(d.keys()):
            if j in i.split():
                l1.append(i)
                break
    final_list_sentences_count.append(len(l1))
    return l1


def word_counter(text):
    l = text.split()
    total_word_count.append(len(l))                   # counts total number of words
    wl = [i for i in l if i in list(d.keys())]        # counts words from dictionary
    final_list_word_count.append(len(wl))

# writes extracted sentences to docx files
def write_docx(l, filename):
    doc = docx.Document()
    doc.add_heading(filename, 0)
    for i in l:
        i = i.replace('\n', ' ')
        i = i.replace('\t', ' ')
        doc.add_paragraph(i)
        doc.add_paragraph('________________________________________________________')
    doc.save(out_path + filename + ".docx")

# extracts text from docx files
def getText(file):
    text = docxpy.process(file)
    return text


try:
    for i, doc in enumerate(glob.iglob(path + "*.docx")):
        filename = doc.split('\\')[-1][0:-5]
        path1 = os.path.abspath(doc)
        txt = getText(path1)
        ftxt = cleaner(txt)
        word_counter(ftxt)
        number_counter(ftxt)
        l = tokenize.sent_tokenize(ftxt)
        l1 = sentence_counter(l)
        final_list_files.append(filename)
        write_docx(l1, filename)



except:
    print(filename)

# stores all lists in dataframe
df = pd.DataFrame.from_dict({'Filename': final_list_files, 'Sentence count': final_list_sentences_count,
                             'Total sentences': total_sentences_count, 'Word count': final_list_word_count,
                             'Total words': total_word_count, 'Number count': final_list_number_count,
                             'Total numbers': total_number_count})
# creates excel file and adds the data in data frame to the file
df.to_excel(out_path + 'vcount.xlsx', header=True, index=False)
